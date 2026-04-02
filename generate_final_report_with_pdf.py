#!/usr/bin/env python3
"""
Convert ML_Report.pdf pages to images and embed in DOCX report using PyMuPDF.
"""
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def shade_cell(cell, fill):
    """Shade cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Convert PDF to images using PyMuPDF
pdf_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/ML_Report.pdf'
output_dir = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/pdf_pages'
os.makedirs(output_dir, exist_ok=True)

print("Converting PDF to images using PyMuPDF...")
doc_pdf = fitz.open(pdf_path)
num_pages = len(doc_pdf)
print(f"✓ PDF has {num_pages} pages")

image_paths = []
for page_num in range(num_pages):
    page = doc_pdf[page_num]
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for quality
    img_path = os.path.join(output_dir, f'page_{page_num + 1:02d}.png')
    pix.save(img_path)
    image_paths.append(img_path)
    print(f"✓ Converted page {page_num + 1}")

doc_pdf.close()

# Now create updated DOCX
doc = Document()

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('AI Resume & Job Match Analyzer')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Short Project Report')
subtitle_run.font.size = Pt(18)
subtitle_run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Authors
authors_heading = doc.add_paragraph()
authors_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_heading_run = authors_heading.add_run('Submitted By:')
authors_heading_run.font.size = Pt(14)
authors_heading_run.font.bold = True

author1 = doc.add_paragraph()
author1.alignment = WD_ALIGN_PARAGRAPH.CENTER
author1_run = author1.add_run('Vaibhav Krishali\n23FE10CSE00301')
author1_run.font.size = Pt(12)

doc.add_paragraph()

author2 = doc.add_paragraph()
author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
author2_run = author2.add_run('Farhat Ansari\n23FE10CSE00151')
author2_run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Course and Date
course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_run = course.add_run('Emerging Technologies\nDepartment of Computer Science and Engineering')
course_run.font.size = Pt(11)

doc.add_paragraph()

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date.add_run('Date: April 2, 2026')
date_run.font.size = Pt(11)

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Report (from LaTeX Compilation)',
    '2. Application Screenshots',
    '3. Conclusion & References'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Project Report Pages from LaTeX
doc.add_heading('1. Project Report (Compiled from LaTeX)', level=1)
doc.add_paragraph('The following pages represent the compiled LaTeX report from Overleaf, containing the detailed project documentation, methodology, results, and analysis.')

doc.add_page_break()

# Add all relevant PDF pages (usually first few pages)
# We'll add pages 1 through min(5, len(images)) to keep report focused
max_pages = min(5, len(image_paths))  # First 5 pages or all if fewer
for i in range(max_pages):
    if i < len(image_paths):
        img_path = image_paths[i]
        try:
            print(f"Embedding page {i+1}...")
            doc.add_picture(img_path, width=Inches(6.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page label
            label = doc.add_paragraph(f'Page {i+1} of compiled report')
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in label.runs:
                run.font.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add page break between embedded pages if not the last one
            if i < max_pages - 1:
                doc.add_page_break()
        except Exception as e:
            print(f"Error embedding page {i+1}: {e}")

doc.add_page_break()

# Application Screenshots
doc.add_heading('2. Application Screenshots & Workflow', level=1)
doc.add_paragraph('The following section shows the live application interface and workflow for resume analysis:')

base_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/report_assets'
if os.path.exists(base_path):
    screenshots = sorted([f for f in os.listdir(base_path) if f.endswith('.png')])
    for i, screenshot in enumerate(screenshots[:4], 1):  # Embed first 4 screenshots
        img_path = os.path.join(base_path, screenshot)
        if os.path.exists(img_path):
            try:
                print(f"Embedding screenshot {i}...")
                doc.add_paragraph()
                doc.add_picture(img_path, width=Inches(6.0))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                desc = doc.add_paragraph(f'Figure {i}: {screenshot.replace("_", " ").replace(".png", "")}')
                desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in desc.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)
            except Exception as e:
                print(f"Error embedding {screenshot}: {e}")

doc.add_page_break()

# Conclusion
doc.add_heading('3. Conclusion & Project Summary', level=1)
doc.add_paragraph("""This AI Resume & Job Match Analyzer successfully demonstrates a comprehensive solution for resume optimization and role-specific analysis. The project combines:

- Automated CV scanning with ATS-style scoring
- Role-specific skill matching (ML Engineer example shown)
- Resume upgrade recommendations with before/after comparison
- Multiple analysis modes (OpenAI, Groq, Local LLM, Heuristic fallback)
- Professional export formats (LaTeX, PDF, Web interface)

The compiled report above details the technical architecture, methodology, implementation, and results. The application screenshots demonstrate the practical usability and effectiveness of the system in real-world scenarios.""")

doc.add_paragraph()
doc.add_heading('Key Features Implemented:', level=2)
features = [
    'CV Readiness Scoring (0-100 scale with ATS assessment)',
    'Role-Specific Skill Gap Analysis',
    'Intelligent Resume Enhancement with Actionable Suggestions',
    'Multiple LLM Backend Support (Graceful Fallback Chain)',
    'LaTeX Resume Generation & PDF Compilation',
    'Production-Ready FastAPI Architecture',
    'Responsive Web Interface with Real-Time Analysis'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Technical Stack:', level=2)
doc.add_paragraph('Backend: Python 3.13, FastAPI, Uvicorn')
doc.add_paragraph('Frontend: HTML5, CSS3, Vanilla JavaScript')
doc.add_paragraph('PDF Processing: pypdf, pdflatex')
doc.add_paragraph('LLM Integration: OpenAI API, Groq API, Ollama')
doc.add_paragraph('Deployment: Docker, Docker Compose')

# Save final document
output_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/Project_Report_IEEE_Final.docx'
doc.save(output_path)

print(f'\n✅ Final DOCX Report Generated!')
print(f'📄 File: {output_path}')
print(f'\n📊 Report Contents:')
print(f'   ✓ Professional title page with both authors')
print(f'   ✓ Table of Contents')
print(f'   ✓ {max_pages} pages from compiled LaTeX PDF (ML_Report.pdf)')
print(f'   ✓ Application workflow screenshots')
print(f'   ✓ Conclusion with key features and technical stack')
print(f'   ✓ Professional formatting and spacing')
