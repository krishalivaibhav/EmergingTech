#!/usr/bin/env python3
"""
Generate final DOCX report with embedded LaTeX PDF pages only (no app screenshots).
"""
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def shade_cell(cell, fill):
    """Shade cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Convert PDF to images using PyMuPDF
pdf_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/ML_Report.pdf'
output_dir = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/pdf_pages'
os.makedirs(output_dir, exist_ok=True)

print("Converting PDF to images using PyMuPDF...")
doc_pdf = fitz.open(pdf_path)
num_pages = len(doc_pdf)
print(f"✓ PDF has {num_pages} pages")

image_paths = []
for page_num in range(num_pages):
    page = doc_pdf[page_num]
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for quality
    img_path = os.path.join(output_dir, f'page_{page_num + 1:02d}.png')
    pix.save(img_path)
    image_paths.append(img_path)
    print(f"✓ Converted page {page_num + 1}")

doc_pdf.close()

# Now create updated DOCX (without app screenshots)
doc = Document()

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('AI Resume & Job Match Analyzer')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Short Project Report')
subtitle_run.font.size = Pt(18)
subtitle_run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Authors
authors_heading = doc.add_paragraph()
authors_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_heading_run = authors_heading.add_run('Submitted By:')
authors_heading_run.font.size = Pt(14)
authors_heading_run.font.bold = True

author1 = doc.add_paragraph()
author1.alignment = WD_ALIGN_PARAGRAPH.CENTER
author1_run = author1.add_run('Vaibhav Krishali\n23FE10CSE00301')
author1_run.font.size = Pt(12)

doc.add_paragraph()

author2 = doc.add_paragraph()
author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
author2_run = author2.add_run('Farhat Ansari\n23FE10CSE00151')
author2_run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Course and Date
course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_run = course.add_run('Emerging Technologies\nDepartment of Computer Science and Engineering')
course_run.font.size = Pt(11)

doc.add_paragraph()

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date.add_run('Date: April 2, 2026')
date_run.font.size = Pt(11)

doc.add_page_break()

# Project Report Pages from LaTeX
doc.add_heading('Project Report - Compiled from LaTeX', level=1)
doc.add_paragraph('The following pages contain the complete project report compiled from the LaTeX source in Overleaf, including abstract, introduction, literature review, methodology, implementation details, results, and conclusions.')

doc.add_page_break()

# Add all PDF pages
for i in range(len(image_paths)):
    img_path = image_paths[i]
    try:
        print(f"Embedding page {i+1}...")
        doc.add_picture(img_path, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page label
        label = doc.add_paragraph(f'Page {i+1} of Report')
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in label.runs:
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page break between embedded pages if not the last one
        if i < len(image_paths) - 1:
            doc.add_page_break()
    except Exception as e:
        print(f"Error embedding page {i+1}: {e}")

# Save final document
output_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/Project_Report_IEEE_Final.docx'
doc.save(output_path)

print(f'\n✅ Final DOCX Report Updated!')
print(f'📄 File: {output_path}')
print(f'\n📊 Report Contents (Clean Version):')
print(f'   ✓ Professional title page with both authors')
print(f'   ✓ {len(image_paths)} pages from compiled LaTeX PDF')
print(f'   ✓ No application screenshots (removed)')
print(f'   ✓ Professional formatting and spacing')
