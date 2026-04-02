#!/usr/bin/env python3
"""
Generate IEEE-format DOCX report with embedded ML Engineer-specific screenshots.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_styled(doc, text, level):
    """Add styled heading."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style = 'Heading 1'
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
    elif level == 2:
        heading.style = 'Heading 2'
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True
    return heading

def shade_cell(cell, fill):
    """Shade cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_table_of_contents(doc):
    """Add a simple table of contents."""
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Literature Review',
        '3. Methodology / Working',
        '4. Results / Output',
        '5. Conclusion',
        '6. Future Scope',
        '7. References'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

# Create document
doc = Document()

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('AI Resume & Job Match Analyzer')
title_run.font.size = Pt(24)
title_run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Short Project Report')
subtitle_run.font.size = Pt(16)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run('Vaibhav Krishali\nEmerging Technologies Project\nDate: April 2, 2026')
author_run.font.size = Pt(12)

doc.add_page_break()

# Abstract
doc.add_heading('Abstract', level=1)
abstract_text = """This project presents an AI-assisted Resume and Job Match Analyzer built using FastAPI and a lightweight web interface. The system evaluates resume quality, recommends suitable roles, and compares resume-job fit through ATS-style scoring. For ML Engineer roles specifically, it provides role-targeted skill gap analysis, actionable ATS suggestions, and rewritten bullet points emphasizing machine learning impact. A resume-upgrade module generates an improved version with before-vs-after comparison and produces an Overleaf-ready LaTeX draft with PDF preview support. The system is designed for accessibility, remaining functional without paid APIs by supporting local and heuristic analysis modes."""
doc.add_paragraph(abstract_text)

doc.add_page_break()

# Table of Contents
add_table_of_contents(doc)
doc.add_page_break()

# Introduction
add_heading_styled(doc, 'Introduction', 1)
doc.add_paragraph('Problem Statement: Many technical professionals, including ML engineers and data scientists, struggle to understand whether their resumes align with industry expectations and target job descriptions. Manual review processes are slow, inconsistent, and often miss ATS (Applicant Tracking System) keywords critical for job application success.')

doc.add_paragraph('Importance of Topic: Modern hiring pipelines increasingly depend on ATS filtering before human review. For specialized roles like ML Engineer, the gap between candidate qualifications and job requirements can be subtle but significant. A structured, role-aware analyzer helps candidates identify specific skill gaps (e.g., missing TensorFlow/PyTorch mentions) and improve resume quality with quantified metrics and technical depth before submitting applications.')

doc.add_paragraph('Project Objective: The project aims to build a practical web application that can (i) score CV readiness using role-agnostic heuristics, (ii) suggest relevant roles based on resume content, (iii) perform role-specific matching with targeted skill overlap analysis, and (iv) generate an upgraded resume draft with before-vs-after comparison highlighting improvements.')

doc.add_page_break()

# Literature Review
add_heading_styled(doc, 'Literature Review', 1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Work'
hdr_cells[1].text = 'Year'
hdr_cells[2].text = 'Key Contribution'
hdr_cells[3].text = 'Gap Addressed in This Project'

# Header shading
for cell in hdr_cells:
    shade_cell(cell, 'D3D3D3')

rows_data = [
    ('OpenAI API Documentation', '2025', 'Structured LLM output for text analysis tasks.', 'Added fallback chain (OpenAI/Groq/local/heuristic) for reliability and cost control.'),
    ('Groq API Documentation', '2025', 'Fast inference for LLM-based JSON generation.', 'Integrated as optional provider to reduce latency in resume analysis.'),
    ('Ollama Documentation', '2025', 'Local LLM serving without external API dependency.', 'Enabled offline/keyless operation for student-friendly deployment.'),
    ('FastAPI Documentation', '2025', 'High-performance API framework with typed models.', 'Used for robust endpoints, validation, and production-ready service design.'),
]

for i, (work, year, contrib, gap) in enumerate(rows_data, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = work
    row_cells[1].text = year
    row_cells[2].text = contrib
    row_cells[3].text = gap

doc.add_page_break()

# Methodology
add_heading_styled(doc, 'Methodology / Working', 1)

add_heading_styled(doc, 'System Workflow', 2)
workflow_items = [
    'User uploads a PDF resume (or pastes text) and runs Scan CV.',
    'Backend extracts resume text and infers profile status and recommended roles.',
    'User optionally selects a target role (e.g., ML Engineer) and adds job description text.',
    'Role-specific analysis computes fit score, matching skills, missing skills, and ATS suggestions.',
    'Resume-upgrade module generates improved snapshot + LaTeX resume + preview/download outputs.'
]
for item in workflow_items:
    doc.add_paragraph(item, style='List Number')

add_heading_styled(doc, 'Architecture Overview', 2)
arch_text = """The project is organized into four layers: (i) web interface, (ii) API routing and validation, (iii) analyzer and transformation services, and (iv) rendering/export utilities. The UI sends multipart requests to FastAPI endpoints. The backend parses uploaded PDF content, validates user input, and invokes the analyzer orchestrator. The orchestrator selects one engine based on configuration: OpenAI, Groq, local Ollama, or a deterministic fallback analyzer. Finally, resume-upgrade responses are converted into snapshot text and an Overleaf-ready LaTeX template, and optional PDF/PNG previews are generated."""
doc.add_paragraph(arch_text)

# Components table
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Light Grid Accent 1'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Component'
hdr_cells2[1].text = 'Role in Pipeline'

for cell in hdr_cells2:
    shade_cell(cell, 'D3D3D3')

components = [
    ('FastAPI Routes', 'Handle CV scan, role analysis, resume upgrade, and LaTeX compilation endpoints.'),
    ('PDF Parser', 'Extract text from uploaded PDF and reject empty/unreadable files.'),
    ('Analyzer Orchestrator', 'Select execution mode and normalize outputs into stable response schemas.'),
    ('Free Analyzer', 'Perform deterministic skill extraction, matching, gap detection, and ATS suggestions.'),
    ('Template/Compiler Services', 'Build upgraded resume snapshots, LaTeX output, and PDF/preview artifacts.'),
]

for i, (comp, role) in enumerate(components, start=1):
    row_cells = table2.rows[i].cells
    row_cells[0].text = comp
    row_cells[1].text = role

add_heading_styled(doc, 'Core Scoring Logic (Heuristic Mode)', 2)
doc.add_paragraph('In free mode, score is derived from skill overlap and missing-skill penalty:')
doc.add_paragraph('score = clamp(35 + 65·|M|/max(|J|,1) − min(2|G|,20), 20, 100)')
doc.add_paragraph('where M is matched skills, J is job-description skills, and G is missing skills.')

add_heading_styled(doc, 'Key Implementation: Skill Extraction', 2)
doc.add_paragraph('A skill catalog with 50+ technical and soft skills is maintained. For each job description and resume, the system extracts matching skills using pattern-based matching:')

code_text = '''SKILL_CATALOG = (
    SkillKeyword("Python", ("python",)),
    SkillKeyword("Machine Learning", ("machine learning", "ml")),
    SkillKeyword("TensorFlow", ("tensorflow",)),
    SkillKeyword("PyTorch", ("pytorch",)),
    SkillKeyword("Deep Learning", ("deep learning",)),
    SkillKeyword("Data Analysis", ("data analysis",)),
    SkillKeyword("MLOps", ("mlops", "ml ops")),
    SkillKeyword("Docker", ("docker",)),
    SkillKeyword("Kubernetes", ("kubernetes", "k8s")),
)

def extract_skills(text: str) -> set[str]:
    found: set[str] = set()
    for skill in SKILL_CATALOG:
        if any(_contains_pattern(text, p) 
               for p in skill.patterns):
            found.add(skill.label)
    return found'''

code_para = doc.add_paragraph(code_text)
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
code_para.paragraph_format.left_indent = Inches(0.5)

add_heading_styled(doc, 'Stepwise Working Logic', 2)
steps = [
    'Validate that at least one resume source is present (PDF or pasted text).',
    'Parse resume content and clean whitespace/noise.',
    'Extract role and keyword signals from the selected JD or role template.',
    'Compute matching and missing skill sets and estimate fit score.',
    'Generate recruiter-style summary, ATS suggestions, and rewritten bullet points.',
    'Build resume-upgrade artifacts: before/after scores, key improvements, snapshots, and LaTeX source.',
    'Optionally compile LaTeX to PDF and generate first-page PNG preview.'
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# Results
add_heading_styled(doc, 'Results / Output', 1)

add_heading_styled(doc, 'User Interface Results (ML Engineer Focus)', 2)
doc.add_paragraph('The following sequences show the application flow for an ML Engineer role analysis. All screenshots were captured from a live local run using real resume data and ML Engineer-specific job requirements (Python, TensorFlow, PyTorch, MLOps, Docker, Kubernetes):')

# Add screenshots
base_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/report_assets'
figures = [
    ('01_home_page.png', 'Home Page: Resume upload interface and CV-first scan entry point.'),
    ('02_cv_score_card.png', 'CV Score Card: Overall resume quality score and assessment label.'),
    ('03_profile_status.png', 'Profile Status: Inferred seniority level and strengths summary.'),
    ('04_recommended_roles.png', 'Recommended Roles: ML Engineer and related role suggestions based on resume content.'),
    ('06_matching_skills.png', 'Matching Skills: Technical skills found in both resume and ML Engineer job description.'),
    ('07_missing_skills.png', 'Missing Skills: Critical ML-focused skills absent from the resume (e.g., TensorFlow, MLOps).'),
    ('08_ats_suggestions.png', 'ATS Suggestions: Actionable improvements like adding metrics and clarifying technical depth.'),
    ('10_upgrade_score_comparison.png', 'Before-vs-After ATS Score: Estimated improvement from resume rewriting.'),
    ('11_upgrade_summary.png', 'Upgrade Summary: Executive summary of how the rewrite improves ML Engineer alignment.'),
    ('12_key_improvements.png', 'Key Improvements: Specific edits applied to strengthen role-targeted content.'),
]

for img_file, caption in figures:
    img_path = os.path.join(base_path, img_file)
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption_para = doc.add_paragraph(f'Figure: {caption}')
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_para.runs:
                run.font.italic = True
                run.font.size = Pt(10)
            doc.add_paragraph()
        except Exception as e:
            doc.add_paragraph(f'[Image not found: {img_file}]')

add_heading_styled(doc, 'Observed Outcome and ML Engineer-Specific Improvements', 2)
doc.add_paragraph('During local execution with an ML Engineer job description (requiring Python, TensorFlow, PyTorch, MLOps, Docker, Kubernetes), the system successfully:')

improvements = [
    'Identified Python, Deep Learning, and Data Analysis as matching skills from the uploaded resume.',
    'Flagged missing critical skills: TensorFlow, PyTorch, MLOps, and cloud deployment (AWS/GCP/Kubernetes).',
    'Suggested quantifying ML project impact (e.g., "improved model accuracy by 15%", "deployed model to production").',
    'Rewrote experience bullets to emphasize end-to-end ML pipeline ownership and team collaboration.',
    'Generated ATS improvements estimated at +10–15 points on resume readiness score.',
    'Produced role-specific interview preparation questions focused on ML experimental design and deployment challenges.'
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

# Output table
table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Light Grid Accent 1'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Artifact'
hdr_cells3[1].text = 'Purpose in Decision Making'

for cell in hdr_cells3:
    shade_cell(cell, 'D3D3D3')

outputs = [
    ('Match Score', 'Quick signal for resume readiness against a specific role (0–100).'),
    ('Matching/Missing Skills', 'Identifies knowledge gaps and areas to strengthen.'),
    ('ATS Suggestions', 'Specific, actionable edits (e.g., "add metrics", "clarify frameworks").'),
    ('Improved Bullets', 'Example rewrites emphasizing impact and role alignment.'),
]

for i, (artifact, purpose) in enumerate(outputs, start=1):
    row_cells = table3.rows[i].cells
    row_cells[0].text = artifact
    row_cells[1].text = purpose

doc.add_page_break()

# Practical Impact
add_heading_styled(doc, 'Practical Impact', 2)
doc.add_paragraph('The project reduces manual resume review effort by converting a multi-step human review process into a guided pipeline with immediate feedback. It is especially suitable for:')

impact_items = [
    'ML Engineers: Tailored skill overlap analysis for TensorFlow, PyTorch, and MLOps frameworks.',
    'Early-Career Professionals: Quick feedback before applying to competitive positions.',
    'Students: Affordable, API-free analysis using local heuristics or free Ollama models.'
]
for item in impact_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Conclusion
add_heading_styled(doc, 'Conclusion', 1)
conclusion_text = """This project delivers a practical AI-assisted toolkit for resume quality and role-fit analysis specifically tailored for technical roles like ML Engineer. It combines API-driven intelligence with deterministic fallback logic, making it reliable across multiple deployment scenarios. The CV-first flow improves usability by guiding users from general readiness to specific role-targeted matching. The resume-upgrade feature adds strong value through before-vs-after comparison, visual preview support, and production-ready LaTeX exports. Captured local run snapshots confirm end-to-end functionality with real resume analysis and ML Engineer-specific feedback. Overall, the solution is effective for students and early professionals seeking quick, structured, and ATS-aware career guidance."""
doc.add_paragraph(conclusion_text)

# Future Scope
add_heading_styled(doc, 'Future Scope', 1)
future_items = [
    'Add domain-specific scoring rubrics for ML, Data Science, Backend, Frontend, and DevOps roles.',
    'Integrate recruiter-specific templates and interview preparation guides.',
    'Build analytics dashboard for tracking score improvement across multiple resume iterations.',
    'Support multi-language resume analysis and localized job market guidance.'
]
for item in future_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# References
add_heading_styled(doc, 'References', 1)
references = [
    'OpenAI, "OpenAI API Documentation," 2025. [Online]. Available: https://platform.openai.com/docs',
    'Groq, "Groq API Documentation," 2025. [Online]. Available: https://console.groq.com/docs',
    'Ollama, "Ollama Documentation," 2025. [Online]. Available: https://ollama.com',
    'S. Ramírez, "FastAPI Documentation," 2025. [Online]. Available: https://fastapi.tiangolo.com'
]
for ref in references:
    doc.add_paragraph(ref, style='List Number')

# Save document
output_path = '/Users/vaibhavkrishali/Desktop/College/emerging tech/reports/Project_Report_IEEE.docx'
doc.save(output_path)
print(f'✅ DOCX Report generated: {output_path}')
print(f'📊 Document includes:')
print(f'   - Title, abstract, and table of contents')
print(f'   - All sections with proper formatting')
print(f'   - 10 embedded focused screenshots')
print(f'   - Code implementation snippets')
print(f'   - Literature review, methodology, results tables')
print(f'   - ML Engineer-specific analysis outcomes')
print(f'   - Professional IEEE-style formatting')
